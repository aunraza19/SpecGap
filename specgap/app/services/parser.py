import pdfplumber
import io
import base64
import hashlib
from typing import Tuple, Dict, Any, List
from fastapi import UploadFile
from PIL import Image
import pandas as pd

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from app.services.sanitizer import sanitize_document_text


def compute_file_hash(file_bytes: bytes) -> str:
    
    return hashlib.sha256(file_bytes).hexdigest()

def smart_chunk_text(text: str, max_tokens: int = 8000) -> List[str]:
   
    chars_limit = max_tokens * 3  
    chunks = []
    
    while len(text) > chars_limit:
        split_idx = text.rfind('\n', 0, chars_limit)
        if split_idx == -1:
            split_idx = chars_limit
            
        chunks.append(text[:split_idx])
        text = text[split_idx:]
        
    if text:
        chunks.append(text)
        
    return chunks

def validate_file(file: UploadFile) -> bool:
    """Basic validation"""
    return str(file.filename).lower().endswith(('.pdf', '.txt', '.md', '.docx'))
    
async def classify_document(text: str, filename: str) -> Dict[str, Any]:
    """
    Heuristic-based classification of document type.
    """
    text_lower = text[:5000].lower()
    fname_lower = filename.lower()
    
    doc_type = "unknown"
    confidence = 0.5
    
    if "payment" in text_lower or "invoice" in fname_lower:
        doc_type = "invoice"
        confidence = 0.9
    elif "contract" in text_lower or "agreement" in fname_lower or "license" in text_lower:
        doc_type = "legal_contract"
        confidence = 0.85
    elif "proposal" in fname_lower or "statement of work" in text_lower or "sow" in text_lower:
        doc_type = "consultant_proposal"
        confidence = 0.8
    elif "technical" in text_lower or "specification" in text_lower or "architecture" in text_lower:
        doc_type = "tech_spec"
        confidence = 0.8
        
    agents = []
    if doc_type in ["legal_contract", "consultant_proposal"]:
        agents.append("legal")
        agents.append("financial")
    if doc_type in ["tech_spec", "consultant_proposal"]:
        agents.append("architect")
        agents.append("security")
        
    return {
        "detected_type": doc_type,
        "confidence": confidence,
        "recommended_agents": list(set(agents))
    }

async def extract_text_with_ocr(file_bytes: bytes) -> str:
    """
    OCR fallback for scanned PDFs using Tesseract.
    """
    if not OCR_AVAILABLE:
        return "Error: OCR libraries not installed. Install pytesseract and pdf2image."
    
    text_content = ""
    try:
        images = convert_from_bytes(file_bytes, dpi=300)
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang='eng')
            if page_text.strip():
                text_content += f"--- Page {i + 1} (OCR) ---\n{page_text}\n\n"
        return text_content if text_content.strip() else "Error: OCR could not extract any text."
    except Exception as e:
        return f"Error during OCR: {str(e)}"

async def extract_text_from_pdf(file_bytes: bytes, force_ocr: bool = False) -> str:
    """
    Extracts text AND TABLES from a PDF file stream.
    """
    if force_ocr:
        return await extract_text_with_ocr(file_bytes)
    
    text_content = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_num = i + 1
                page_content = f"\n--- PAGE {page_num} ---\n"
                
                try:
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if not table or not any(row for row in table): continue
                            df = pd.DataFrame(table[1:], columns=table[0]) if len(table) > 1 else pd.DataFrame(table)
                            page_content += f"\n[Table]\n{df.to_markdown(index=False)}\n\n"
                except: pass 
                
                text = page.extract_text()
                if text: page_content += text + "\n"
                text_content += page_content
        
        avg_chars = len(text_content) / max(total_pages, 1)
        if not text_content.strip() or avg_chars < 50:
            if OCR_AVAILABLE:
                ocr = await extract_text_with_ocr(file_bytes)
                if "Error" not in ocr: return ocr
            return "Error: No text found (likely scanned)."
            
        return text_content
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"

async def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text and tables from a DOCX file.
    """
    if not DOCX_AVAILABLE:
        return "Error: DOCX libraries not installed. Install python-docx."

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        parts = []

        # Paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in document.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(cells)
            if table_rows:
                try:
                    df = pd.DataFrame(table_rows[1:], columns=table_rows[0]) if len(table_rows) > 1 else pd.DataFrame(table_rows)
                    parts.append("[Table]\n" + df.to_markdown(index=False))
                except Exception:
                    # Fallback: raw table rows
                    parts.append("[Table]\n" + "\n".join([" | ".join(r) for r in table_rows]))

        text_content = "\n\n".join(parts)
        return text_content if text_content.strip() else "Error: No text found in DOCX."
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"

async def extract_text_from_file(file: UploadFile) -> Tuple[str, Dict]:
    """
    Universal extractor that handles PDF, TXT, MD, etc.
    Returns (text, metadata)
    """
    content = await file.read()
    filename = file.filename.lower()
    
    text = ""
    metadata = {
        "filename": file.filename,
        "size_bytes": len(content),
        "content_type": file.content_type
    }
    
    if filename.endswith(".pdf"):
        text = await extract_text_from_pdf(content)
        metadata["format"] = "pdf"
    elif filename.endswith(".docx"):
        text = await extract_text_from_docx(content)
        metadata["format"] = "docx"
    elif filename.endswith(".txt") or filename.endswith(".md"):
        try:
            text = content.decode("utf-8")
        except:
            text = content.decode("latin-1")
        metadata["format"] = "text"
    else:
        # Fallback for now or error
        text = f"Error: Unsupported file format {filename}. Only PDF, DOCX, TXT, MD supported."
        metadata["format"] = "unknown"

    # Sanitize extracted text to prevent prompt injection (Test Case 5)
    if not text.startswith("Error:"):
        text = sanitize_document_text(text, max_length=500000)

    return text, metadata

def encode_image_for_gemini(image_file: bytes, mime_type: str = "image/png"):
    return {
        "mime_type": mime_type,
        "data": base64.b64encode(image_file).decode("utf-8")
    }
